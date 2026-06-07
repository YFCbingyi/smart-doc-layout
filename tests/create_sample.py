"""生成示例Word文档用于测试。"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


def create_sample_docx(output_path: str = "examples/sample.docx") -> str:
    """创建一个包含多种样式的示例Word文档。"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # 添加标题
    heading = doc.add_heading("Word文档识别与修改测试", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    # 添加一级标题
    doc.add_heading("第一章 概述", level=1)

    # 添加正文段落
    para1 = doc.add_paragraph()
    para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run1 = para1.add_run("本文档用于测试Word文档的识别与修改功能。")
    run1.font.name = "宋体"
    run1.font.size = Pt(12)
    run2 = para1.add_run("其中包含多种样式和格式。")
    run2.font.name = "宋体"
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(255, 0, 0)

    # 添加二级标题
    doc.add_heading("1.1 背景介绍", level=2)

    # 添加正文
    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para2.paragraph_format.first_line_indent = Cm(0.74)
    para2.add_run("在办公自动化领域，Word文档的处理是一个常见需求。").font.size = Pt(12)

    # 添加列表
    doc.add_paragraph("项目A", style="List Bullet")
    doc.add_paragraph("项目B", style="List Bullet")
    doc.add_paragraph("项目C", style="List Bullet")

    # 添加三级标题
    doc.add_heading("1.2 技术路线", level=2)

    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para3.paragraph_format.first_line_indent = Cm(0.74)
    para3.add_run("本项目使用python-docx库进行Word文档的读写操作。").font.size = Pt(12)

    # 添加页眉
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "内部测试文档"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # 添加页脚（页码）
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("第 ")
    run.font.size = Pt(9)
    # 页码字段通过OXML插入
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    footer_para._element.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    footer_para._element.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    footer_para._element.append(fld_char_end)

    run2 = footer_para.add_run(" 页")
    run2.font.size = Pt(9)

    # 保存
    doc.save(output_path)
    print(f"示例文档已创建: {output_path}")
    return output_path


if __name__ == "__main__":
    import os

    os.makedirs("examples", exist_ok=True)
    create_sample_docx()
