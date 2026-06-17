"""Word文档读取识别模块。

读取 .docx 文件，提取段落类型、样式、页眉、页码等信息，
并支持智能识别：主标题/副标题、四级标题（一、/（一）/1./(1)）。

输出格式与修改输入格式完全一致（使用 ModifyInput 模型）。
"""

from __future__ import annotations

import os
import re
from typing import Optional

from docx import Document as DocxDocument
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.section import Section, _Header, _Footer
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from word_processor.models import (
    AlignmentType,
    HeaderFooterInput,
    ModifyInput,
    PageNumberInput,
    ParagraphModify,
    ParagraphType,
    SectionModify,
    TextStyleInput,
)
from word_processor.page_number import has_page_number, get_page_number_format

# ═══════════════════════════════════════════════════
# 四级标题正则模式
# ═══════════════════════════════════════════════════

_PATTERN_LEVEL1 = re.compile(r"^[一二三四五六七八九十百千]+[、]")
_PATTERN_LEVEL2 = re.compile(r"^[（(][一二三四五六七八九十百千]+[）)]")
_PATTERN_LEVEL3 = re.compile(r"^\d+\.[\s]")
_PATTERN_LEVEL4 = re.compile(r"^[（(]\d+[）)]")

_HEADING_PATTERNS: list[tuple[re.Pattern, int]] = [
    (_PATTERN_LEVEL1, 1),
    (_PATTERN_LEVEL2, 2),
    (_PATTERN_LEVEL3, 3),
    (_PATTERN_LEVEL4, 4),
]


def _alignment_to_model(alignment: Optional[WD_ALIGN_PARAGRAPH]) -> AlignmentType:
    """将 python-docx 对齐方式转换为模型枚举。"""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: AlignmentType.left,
        WD_ALIGN_PARAGRAPH.CENTER: AlignmentType.center,
        WD_ALIGN_PARAGRAPH.RIGHT: AlignmentType.right,
        WD_ALIGN_PARAGRAPH.JUSTIFY: AlignmentType.justify,
    }
    return mapping.get(alignment, AlignmentType.unknown)


def _get_font_color(run: Run) -> Optional[str]:
    """获取Run的字体颜色，返回十六进制字符串。"""
    try:
        color = run.font.color
        if color and color.rgb:
            return str(color.rgb)
    except Exception:
        pass
    return None


def _get_paragraph_type(paragraph: Paragraph) -> tuple[ParagraphType, Optional[int]]:
    """根据Word样式名判断段落类型和标题级别（仅基于样式）。"""
    style = paragraph.style
    if style is None or style.name is None:
        return ParagraphType.body, None

    style_name = style.name

    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1])
            return ParagraphType.heading, level
        except (ValueError, IndexError):
            return ParagraphType.heading, 1

    if "List" in style_name or style_name.startswith("List"):
        return ParagraphType.list, None

    return ParagraphType.body, None


def _detect_heading_by_content(para: ParagraphModify) -> tuple[ParagraphType, Optional[int]]:
    """根据段落文本内容智能识别标题类型。"""
    text = (para.text or "").strip()
    if not text:
        return para.type or ParagraphType.body, para.heading_level

    for pattern, level in _HEADING_PATTERNS:
        if pattern.match(text):
            return ParagraphType.heading, level

    return para.type or ParagraphType.body, para.heading_level


def _detect_title_subtitle(paragraphs: list[ParagraphModify]) -> None:
    """智能识别文档正文开头的主标题和副标题（跳过已标记的封面区域）。"""
    # 找到正文起始位置（第一个非封面段落）
    start_idx = 0
    for i, p in enumerate(paragraphs):
        if p.type != ParagraphType.skip:
            start_idx = i
            break

    body_sizes = [
        p.style.font_size
        for p in paragraphs[start_idx:]
        if p.style is not None and p.style.font_size is not None and p.type == ParagraphType.body
    ]
    avg_body_size = sum(body_sizes) / len(body_sizes) if body_sizes else 12.0

    title_candidate: Optional[ParagraphModify] = None
    for p in paragraphs[start_idx:]:
        text = (p.text or "").strip()
        if not text:
            continue
        if p.style is None or p.style.alignment != AlignmentType.center:
            break

        if title_candidate is None:
            p.type = ParagraphType.title
            title_candidate = p
        else:
            same_font = (
                p.style is not None
                and title_candidate.style is not None
                and p.style.font_name == title_candidate.style.font_name
                and p.style.font_size == title_candidate.style.font_size
            )
            p.type = ParagraphType.title if same_font else ParagraphType.subtitle


def _detect_cover_area(paragraphs: list[ParagraphModify]) -> None:
    """识别文档封面区域，标记为 other 类型。

    封面是文档开头的区域，通常包含文档标题、作者、单位等信息，
    这部分不参与样式修改。

    识别逻辑：从文档开头扫描，连续满足以下条件之一即为封面：
    - 空白段落（用于间距）
    - 居中对齐的段落（封面标题/作者）
    - 右对齐的段落（封面单位/日期，需同时满足：
      (a) 前面出现过居中内容；
      (b) 右对齐块之后的下一个非空段落是居中段落 → 是封面落款；
      如果右对齐块之后直接跟正文，则视为副标题，不归入封面）
    遇到以下情况时结束封面区域：
    - 标题段落（正文开始）
    - 左对齐或两端对齐的非空段落（正文内容）
    - 右对齐段落但前面没有居中内容（可能是正文右对齐段落）
    - 在已识别出右对齐封面内容后，后续居中的段落视为正文标题
    """
    if not paragraphs:
        return

    first_content_idx: Optional[int] = None
    saw_centered = False
    saw_right_aligned_in_cover = False

    for i, p in enumerate(paragraphs):
        text = (p.text or "").strip()

        # 空白段落属于封面（用于间距）
        if not text:
            continue

        # 标题段落表示正文开始
        if p.type == ParagraphType.heading:
            first_content_idx = i
            break

        alignment = p.style.alignment if p.style is not None else None

        # 居中对齐的段落
        if alignment == AlignmentType.center:
            if saw_right_aligned_in_cover:
                # 在右对齐封面内容（单位/日期）之后出现居中内容 → 正文标题
                first_content_idx = i
                break
            saw_centered = True
            continue  # 居中属于封面

        # 右对齐的段落
        if alignment == AlignmentType.right:
            if not saw_centered:
                # 前面无居中内容 → 正文中的右对齐段落
                first_content_idx = i
                break
            # 前面有居中内容：检查这是封面落款还是副标题
            # 跳过所有连续右对齐段落，看后面跟什么
            j = i + 1
            while j < len(paragraphs):
                t = (paragraphs[j].text or "").strip()
                a = paragraphs[j].style.alignment if paragraphs[j].style else None
                if t and a != AlignmentType.right:
                    break
                j += 1
            # 找右对齐块之后的下一个非空段落
            next_non_empty = None
            for k in range(j, len(paragraphs)):
                if (paragraphs[k].text or "").strip():
                    next_non_empty = paragraphs[k]
                    break
            if next_non_empty is not None:
                na = next_non_empty.style.alignment if next_non_empty.style else None
                if na != AlignmentType.center and next_non_empty.type != ParagraphType.heading:
                    # 右对齐块后直接跟正文 → 这是副标题，不是封面落款
                    # 整个文档无封面区域
                    first_content_idx = 0
                    break
            # 右对齐块后跟居中内容（正文标题）→ 属于封面落款
            saw_right_aligned_in_cover = True
            continue

        # 左对齐/两端对齐/其他 → 正文开始
        first_content_idx = i
        break

    if first_content_idx is None:
        first_content_idx = len(paragraphs)

    if first_content_idx > 0:
        for p in paragraphs[:first_content_idx]:
            p.type = ParagraphType.skip
            p.heading_level = None


def _detect_signature_area(paragraphs: list[ParagraphModify]) -> None:
    """识别文档末尾的落款区域（单位名称和日期），标记为 other 类型。

    落款特征：文档末尾连续右对齐的非空段落，通常为 1-3 行，
    例如最后两行分别是单位名称和日期。

    识别逻辑：从文档末尾向前扫描最多 10 个段落，收集连续右对齐
    的非空段落及其前面的空白行，标记为 other。
    """
    if not paragraphs:
        return

    n = len(paragraphs)
    scan_start = max(0, n - 10)
    sig_start = n  # 落款起始索引，默认无落款

    found_signature = False
    for i in range(n - 1, scan_start - 1, -1):
        p = paragraphs[i]
        text = (p.text or "").strip()

        if found_signature:
            # 已在落款区域中：空白行或右对齐段落继续扩展区域
            if not text:
                sig_start = i
                continue
            if p.style is not None and p.style.alignment == AlignmentType.right:
                sig_start = i
                continue
            break  # 遇到非空白、非右对齐段落，落款结束
        else:
            # 尚未找到落款
            if not text:
                continue  # 跳过末尾空白行
            if p.style is not None and p.style.alignment == AlignmentType.right:
                found_signature = True
                sig_start = i
                continue
            break  # 非右对齐，不是落款

    if sig_start < n:
        for p in paragraphs[sig_start:]:
            p.type = ParagraphType.skip
            p.heading_level = None


def _extract_paragraph_style(paragraph: Paragraph) -> TextStyleInput:
    """提取段落样式。"""
    pf = paragraph.paragraph_format
    alignment = paragraph.alignment

    font_name = None
    font_size = None
    bold = None
    color = None
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font = first_run.font
        font_name = font.name if font.name else None
        font_size = font.size.pt if font.size else None
        bold = font.bold if font.bold is not None else None
        color = _get_font_color(first_run)

    # 规范化行距：将 OXML EMU 值转换为磅值
    raw_ls = pf.line_spacing
    line_spacing = None
    if raw_ls is not None:
        if isinstance(raw_ls, (int, float)) and raw_ls > 1000:
            line_spacing = round(raw_ls / 12700, 1)  # EMU → 磅
        else:
            line_spacing = raw_ls  # 已经是磅值或倍数

    return TextStyleInput(
        font_name=font_name,
        font_size=font_size,
        bold=bold,
        color=color,
        alignment=_alignment_to_model(alignment),
        line_spacing=line_spacing,
        space_before=pf.space_before.pt if pf.space_before else None,
        space_after=pf.space_after.pt if pf.space_after else None,
        first_line_indent=pf.first_line_indent.pt if pf.first_line_indent else None,
    )


def _extract_paragraph_data(index: int, paragraph: Paragraph) -> ParagraphModify:
    """提取单个段落的数据（基于样式）。"""
    para_type, heading_level = _get_paragraph_type(paragraph)

    return ParagraphModify(
        index=index,
        type=para_type,
        heading_level=heading_level,
        text=paragraph.text,
        style=_extract_paragraph_style(paragraph),
    )


def _extract_header_footer(
    hf_obj: Optional[_Header | _Footer],
) -> Optional[HeaderFooterInput]:
    """提取页眉或页脚的数据，返回 HeaderFooterInput。"""
    if hf_obj is None:
        return None

    try:
        full_text = ""
        content_style: Optional[TextStyleInput] = None
        pn_config: Optional[PageNumberInput] = None

        for i, para in enumerate(hf_obj.paragraphs):
            if para.text:
                if full_text:
                    full_text += "\n"
                full_text += para.text

            # 提取样式（取第一个段落）
            if i == 0:
                content_style = _extract_paragraph_style(para)

            # 检测页码
            if pn_config is None and has_page_number(para):
                fmt = get_page_number_format(para)
                if fmt == "PAGE":
                    fmt = "{PAGE}"
                pn_config = PageNumberInput(enabled=True, format=fmt or "{PAGE}")

        return HeaderFooterInput(
            content=full_text if full_text else None,
            style=content_style,
            page_number=pn_config,
        )
    except Exception:
        return None


def _extract_section_data(index: int, section: Section) -> SectionModify:
    """提取节(Section)的数据。"""
    header = _extract_header_footer(section.header)
    footer = _extract_header_footer(section.footer)

    # 提取偶数页页眉/页脚
    even_header = None
    even_footer = None
    odd_even_enabled = False

    try:
        odd_even_enabled = section.different_first_page_header_footer or section.even_page_header.is_linked_to_previous is not None
        # 实际检查偶数页页眉和页脚是否有内容
        even_header = _extract_header_footer(section.even_page_header)
        even_footer = _extract_header_footer(section.even_page_footer)
        if even_header is not None or even_footer is not None:
            odd_even_enabled = True
    except Exception:
        pass

    # 读取页脚距离
    footer_distance_cm = None
    try:
        if section.footer_distance is not None:
            footer_distance_cm = round(section.footer_distance / 360000, 2)
    except Exception:
        pass

    return SectionModify(
        index=index,
        header=header,
        footer=footer,
        even_page_header=even_header,
        even_page_footer=even_footer,
        enable_odd_even=odd_even_enabled,
        footer_distance_cm=footer_distance_cm,
    )


def read_document(file_path: str, smart: bool = True) -> ModifyInput:
    """读取Word文档，提取结构化数据。

    处理流程：
      1. 按Word样式提取基本段落数据
      2. 智能识别四级标题（一、/（一）/1./(1)）(smart=True时)
      3. 智能识别主标题/副标题 (smart=True时)

    Args:
        file_path: .docx文件路径。
        smart: 是否启用智能识别（识别主副标题、四级标题内容）。
               默认True；设为False时仅基于Word实际样式识别。

    Returns:
        ModifyInput: 统一数据模型，可直接编辑作为修改输入。

    Raises:
        FileNotFoundError: 文件不存在。
        ValueError: 文件格式不支持。
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    if not file_path.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 格式的Word文档")

    doc: DocumentType = DocxDocument(file_path)

    # 提取节数据
    sections = [_extract_section_data(i, section) for i, section in enumerate(doc.sections)]

    # ── 第1步：按样式提取段落数据 ──
    paragraphs = [_extract_paragraph_data(i, para) for i, para in enumerate(doc.paragraphs)]

    if smart:
        # ── 第2步：智能识别四级标题（按内容） ──
        for p in paragraphs:
            if p.type == ParagraphType.body:
                new_type, new_level = _detect_heading_by_content(p)
                if new_type == ParagraphType.heading:
                    p.type = new_type
                    p.heading_level = new_level

        # ── 第3步：智能识别封面区域（标记为skip，不参与样式修改） ──
        _detect_cover_area(paragraphs)

        # ── 第4步：智能识别主标题/副标题（正文开头，跳过封面） ──
        _detect_title_subtitle(paragraphs)

        # ── 第5步：智能识别文末落款（单位名称和日期，标记为skip） ──
        _detect_signature_area(paragraphs)

    return ModifyInput(
        sections=sections,
        paragraphs=paragraphs,
    )
