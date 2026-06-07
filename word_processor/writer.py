"""Word文档修改写入模块。

根据JSON格式的样式和内容数据，修改Word文档的段落样式、内容、页眉、页码。
"""

from __future__ import annotations

import os
from typing import Optional

from docx import Document as DocxDocument
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from word_processor.models import (
    AlignmentType,
    HeaderFooterInput,
    ModifyInput,
    ParagraphModify,
    ParagraphType,
    SectionModify,
    TextStyleInput,
)
from word_processor.page_number import add_page_number, has_page_number, remove_page_number


def _parse_color(hex_color: Optional[str]) -> Optional[RGBColor]:
    """将十六进制颜色字符串转换为RGBColor对象。"""
    if hex_color is None:
        return None
    try:
        hex_str = hex_color.lstrip("#")
        if len(hex_str) == 6:
            return RGBColor(
                int(hex_str[0:2], 16),
                int(hex_str[2:4], 16),
                int(hex_str[4:6], 16),
            )
    except (ValueError, AttributeError):
        pass
    return None


def _alignment_to_docx(alignment: Optional[AlignmentType]) -> Optional[WD_ALIGN_PARAGRAPH]:
    """将模型对齐方式转换为 python-docx 对齐方式。"""
    mapping = {
        AlignmentType.left: WD_ALIGN_PARAGRAPH.LEFT,
        AlignmentType.center: WD_ALIGN_PARAGRAPH.CENTER,
        AlignmentType.right: WD_ALIGN_PARAGRAPH.RIGHT,
        AlignmentType.justify: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(alignment)  # type: ignore[arg-type]


def _apply_style_to_run(run: Run, style: TextStyleInput) -> None:
    """将样式应用到Run对象。"""
    font = run.font

    if style.font_name is not None:
        font.name = style.font_name
        # 同时设置东亚字体（对中文有效）
        rpr = run._element.get_or_add_rPr()
        ea_font = rpr.find(qn("w:rFonts"))
        if ea_font is None:
            ea_font = __import__("docx.oxml", fromlist=["OxmlElement"]).OxmlElement("w:rFonts")
            rpr.insert(0, ea_font)
        if style.font_name:
            ea_font.set(qn("w:eastAsia"), style.font_name)

    if style.font_size is not None:
        font.size = Pt(style.font_size)

    if style.bold is not None:
        font.bold = style.bold

    if style.italic is not None:
        font.italic = style.italic

    if style.underline is not None:
        font.underline = style.underline

    color = _parse_color(style.color)
    if color is not None:
        font.color.rgb = color


def _apply_style_to_paragraph(paragraph: Paragraph, style: TextStyleInput) -> None:
    """将样式应用到段落（段落级格式）。"""
    pf = paragraph.paragraph_format

    if style.alignment is not None:
        paragraph.alignment = _alignment_to_docx(style.alignment)

    if style.line_spacing is not None:
        # 如果行距值 > 10，视为磅值 → 用 Pt() 包装
        if isinstance(style.line_spacing, (int, float)) and style.line_spacing > 10:
            pf.line_spacing = Pt(style.line_spacing)
        else:
            pf.line_spacing = style.line_spacing

    if style.space_before is not None:
        pf.space_before = Pt(style.space_before)

    if style.space_after is not None:
        pf.space_after = Pt(style.space_after)

    if style.first_line_indent is not None:
        pf.first_line_indent = Pt(style.first_line_indent)


def _set_paragraph_text(paragraph: Paragraph, text: str, style: Optional[TextStyleInput] = None) -> None:
    """设置段落文本内容，可同时应用样式。

    清除段落原有所有run，添加一个新的run并设置文本和样式。
    """
    # 清除所有现有run
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    # 添加新run
    new_run = paragraph.add_run(text)

    if style:
        _apply_style_to_run(new_run, style)

    # 确保段落有文本（即使为空）
    if not text and not paragraph.runs:
        paragraph.add_run("")


def _set_heading_level(paragraph: Paragraph, level: int) -> None:
    """设置段落为标题级别。"""
    style_name = f"Heading {level}"
    try:
        paragraph.style = paragraph.part.document.styles[style_name]
    except KeyError:
        # 如果样式不存在，尝试使用默认样式
        pass


def _modify_section_header_footer(
    doc: DocumentType,
    section_index: int,
    is_header: bool,
    modify_data: HeaderFooterInput,
) -> None:
    """修改指定节的页眉或页脚。"""
    section = doc.sections[section_index]
    hf_obj = section.header if is_header else section.footer

    # 确保不linked to previous，才能独立修改
    if hf_obj.is_linked_to_previous:
        hf_obj.is_linked_to_previous = False

    # 修改内容
    if modify_data.content is not None:
        if hf_obj.paragraphs:
            first_para = hf_obj.paragraphs[0]
            _set_paragraph_text(first_para, modify_data.content, modify_data.style)

            # 如果还有更多段落，清除它们
            for extra_para in hf_obj.paragraphs[1:]:
                p_element = extra_para._element
                p_element.getparent().remove(p_element)
        else:
            new_para = hf_obj.add_paragraph(modify_data.content)
            if modify_data.style:
                _apply_style_to_run(new_para.runs[0], modify_data.style)

    # 应用段落级别样式
    if modify_data.style and hf_obj.paragraphs:
        _apply_style_to_paragraph(hf_obj.paragraphs[0], modify_data.style)

    # 处理页码
    if modify_data.page_number is not None:
        pn_config = modify_data.page_number
        first_para = hf_obj.paragraphs[0] if hf_obj.paragraphs else hf_obj.add_paragraph("")

        if pn_config.enabled:
            # 添加页码
            if has_page_number(first_para):
                remove_page_number(first_para)
            add_page_number(first_para, pn_config.format)
        else:
            # 移除页码
            if has_page_number(first_para):
                if first_para.text.strip():
                    remove_page_number(first_para)
                else:
                    # 如果段落只有页码，直接设空
                    remove_page_number(first_para)


def _modify_even_page_footer(
    doc: DocumentType,
    section_index: int,
    modify_data: HeaderFooterInput,
) -> None:
    """修改指定节的偶数页页脚。"""
    section = doc.sections[section_index]

    # 获取偶数页页脚
    even_footer = section.even_page_footer

    # 确保不 linked to previous
    if even_footer.is_linked_to_previous:
        even_footer.is_linked_to_previous = False

    # 修改内容
    if modify_data.content is not None:
        if even_footer.paragraphs:
            first_para = even_footer.paragraphs[0]
            _set_paragraph_text(first_para, modify_data.content, modify_data.style)
            for extra_para in even_footer.paragraphs[1:]:
                p_element = extra_para._element
                p_element.getparent().remove(p_element)
        else:
            new_para = even_footer.add_paragraph(modify_data.content)
            if modify_data.style:
                _apply_style_to_run(new_para.runs[0], modify_data.style)

    # 应用段落样式
    if modify_data.style and even_footer.paragraphs:
        _apply_style_to_paragraph(even_footer.paragraphs[0], modify_data.style)

    # 处理页码
    if modify_data.page_number is not None:
        pn_config = modify_data.page_number
        first_para = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph("")

        if pn_config.enabled:
            if has_page_number(first_para):
                remove_page_number(first_para)
            add_page_number(first_para, pn_config.format)
        else:
            if has_page_number(first_para):
                remove_page_number(first_para)


def _modify_even_page_header(
    doc: DocumentType,
    section_index: int,
    modify_data: HeaderFooterInput,
) -> None:
    """修改指定节的偶数页页眉。"""
    section = doc.sections[section_index]

    # 获取偶数页页眉
    even_header = section.even_page_header

    # 确保不 linked to previous
    if even_header.is_linked_to_previous:
        even_header.is_linked_to_previous = False

    # 修改内容
    if modify_data.content is not None:
        if even_header.paragraphs:
            first_para = even_header.paragraphs[0]
            _set_paragraph_text(first_para, modify_data.content, modify_data.style)
            for extra_para in even_header.paragraphs[1:]:
                p_element = extra_para._element
                p_element.getparent().remove(p_element)
        else:
            new_para = even_header.add_paragraph(modify_data.content)
            if modify_data.style:
                _apply_style_to_run(new_para.runs[0], modify_data.style)

    # 应用段落样式
    if modify_data.style and even_header.paragraphs:
        _apply_style_to_paragraph(even_header.paragraphs[0], modify_data.style)


def _modify_paragraph(doc: DocumentType, modify_data: ParagraphModify) -> None:
    """修改指定段落的内容和样式。"""
    if modify_data.index >= len(doc.paragraphs):
        return

    paragraph = doc.paragraphs[modify_data.index]

    # 修改标题级别
    if modify_data.heading_level is not None:
        _set_heading_level(paragraph, modify_data.heading_level)

    # 修改文本内容
    if modify_data.text is not None:
        _set_paragraph_text(paragraph, modify_data.text, modify_data.style)

    # 应用段落级别样式（即使文本内容未修改）
    if modify_data.style:
        # 如果有runs但文本未修改，对第一个run应用字体样式
        if modify_data.text is None and paragraph.runs:
            _apply_style_to_run(paragraph.runs[0], modify_data.style)
        _apply_style_to_paragraph(paragraph, modify_data.style)


def modify_document(
    file_path: str,
    output_path: str,
    modify_data: ModifyInput,
) -> None:
    """根据输入数据修改Word文档。

    Args:
        file_path: 源文档路径。
        output_path: 输出文档路径。
        modify_data: 修改参数模型，包含要修改的节和段落信息。

    Raises:
        FileNotFoundError: 文件不存在。
        ValueError: 文件格式不支持。
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    if not file_path.lower().endswith(".docx"):
        raise ValueError("仅支持 .docx 格式的Word文档")

    doc: DocumentType = DocxDocument(file_path)

    # 1. 修改节（页眉/页脚）
    for section_mod in modify_data.sections:
        if section_mod.index >= len(doc.sections):
            continue

        # 启用/禁用奇偶分页
        if section_mod.enable_odd_even:
            doc.settings.odd_and_even_pages_header_footer = True
        else:
            doc.settings.odd_and_even_pages_header_footer = False
            # 清除已有的偶数页页眉/页脚内容（如果存在）
            section = doc.sections[section_mod.index]
            try:
                even_header = section.even_page_header
                if even_header and even_header.paragraphs:
                    for p in list(even_header.paragraphs):
                        if p.runs:
                            for r in p.runs:
                                r.text = ""
                        p._element.getparent().remove(p._element)
            except Exception:
                pass
            try:
                even_footer = section.even_page_footer
                if even_footer and even_footer.paragraphs:
                    for p in list(even_footer.paragraphs):
                        if p.runs:
                            for r in p.runs:
                                r.text = ""
                        p._element.getparent().remove(p._element)
            except Exception:
                pass

        if section_mod.header is not None:
            _modify_section_header_footer(
                doc, section_mod.index, is_header=True, modify_data=section_mod.header
            )

        if section_mod.footer is not None:
            _modify_section_header_footer(
                doc, section_mod.index, is_header=False, modify_data=section_mod.footer
            )

        # 处理偶数页页脚（仅启用奇偶分页时）
        if section_mod.enable_odd_even and section_mod.even_page_footer is not None:
            _modify_even_page_footer(doc, section_mod.index, section_mod.even_page_footer)

        # 处理偶数页页眉（仅启用奇偶分页时）
        if section_mod.enable_odd_even and section_mod.even_page_header is not None:
            _modify_even_page_header(doc, section_mod.index, section_mod.even_page_header)

    # 2. 修改段落
    for para_mod in modify_data.paragraphs:
        _modify_paragraph(doc, para_mod)

    # 保存
    doc.save(output_path)
